from datetime import datetime
from pathlib import Path

import aiofiles

from app.connectors.base import BaseConnector, RawDocument


class LocalConnector(BaseConnector):
    """本地文件连接器，用于处理上传的文件"""

    def __init__(self, upload_dir: str):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)

    async def fetch_documents(self, since: datetime | None = None) -> list[RawDocument]:
        results = []
        for f in self.upload_dir.iterdir():
            if f.is_file():
                content = await self._read_file(f)
                if content:
                    results.append(
                        RawDocument(
                            external_id=f.name,
                            title=f.stem,
                            content=content,
                            metadata={"path": str(f), "size": f.stat().st_size},
                        )
                    )
        return results

    async def fetch_document(self, doc_id: str) -> RawDocument:
        f = self.upload_dir / doc_id
        content = await self._read_file(f)
        return RawDocument(
            external_id=doc_id,
            title=f.stem,
            content=content,
            metadata={"path": str(f), "size": f.stat().st_size},
        )

    async def save_file(self, filename: str, data: bytes) -> str:
        """保存上传的文件，返回文件名"""
        filepath = self.upload_dir / filename
        async with aiofiles.open(filepath, "wb") as f:
            await f.write(data)
        return filename

    @staticmethod
    async def _read_file(path: Path) -> str | None:
        suffix = path.suffix.lower()
        if suffix in (".txt", ".md", ".markdown"):
            async with aiofiles.open(path, "r", encoding="utf-8", errors="ignore") as f:
                return await f.read()
        elif suffix == ".pdf":
            return await _extract_pdf(path)
        elif suffix == ".docx":
            return await _extract_docx(path)
        elif suffix == ".xlsx":
            return await _extract_xlsx(path)
        return None


async def _extract_pdf(path: Path) -> str:
    """提取PDF文本内容"""
    import subprocess

    try:
        result = subprocess.run(
            ["pdftotext", "-layout", str(path), "-"], capture_output=True, text=True, timeout=30
        )
        return result.stdout if result.returncode == 0 else None
    except Exception:
        return None


async def _extract_docx(path: Path) -> str:
    """提取DOCX文本内容"""
    from docx import Document as DocxDocument

    doc = DocxDocument(str(path))
    return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())


async def _extract_xlsx(path: Path) -> str:
    """提取XLSX文本内容"""
    from openpyxl import load_workbook

    wb = load_workbook(str(path), read_only=True)
    rows = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                rows.append("\t".join(cells))
    return "\n".join(rows)
