import asyncio
import hashlib
import logging
from pathlib import Path
from uuid import UUID

import aiofiles
from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse
from pydantic import BaseModel
from sqlalchemy import case, delete, func, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.config import settings
from app.core.cache import cache_delete, cache_get, cache_set
from app.core.ratelimit import upload_rate_limit
from app.core.security import get_current_user
from app.database import get_db
from app.models.document import DataSource, Document, DocumentChunk
from app.models.permission import DocumentPermission, SourcePermission
from app.models.user import User
from app.services.audit import log as audit_log
from app.services.webhook import dispatch as webhook_dispatch

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/documents", tags=["文档"])

ALLOWED_EXT = set(settings.ALLOWED_EXTENSIONS.split(","))


def _permission_allows(permission: str, required: str) -> bool:
    """判断权限等级是否满足需求。write 隐含 read。"""
    if permission == "write":
        return True
    if permission == "read" and required == "read":
        return True
    return False


async def _require_doc_permission(
    db: AsyncSession, doc_id: str, user: User, action: str = "查看", required: str = "read"
) -> None:
    """检查用户对文档的权限，无权限时抛出 403

    required: "read" 表示需要读权限，"write" 表示需要写权限
    DocumentPermission 优先；仅当无 DocumentPermission 时才回退到 SourcePermission。
    最后检查知识库成员权限。
    """
    if user.role == "admin":
        return

    # 查 DocumentPermission（优先）
    perm = await db.execute(
        select(DocumentPermission).where(
            DocumentPermission.document_id == doc_id,
            DocumentPermission.user_id == user.id,
        )
    )
    doc_perm = perm.scalar_one_or_none()
    if doc_perm:
        if not _permission_allows(doc_perm.permission, required):
            raise HTTPException(status_code=403, detail=f"无权限{action}该文档")
        return

    # 无 DocumentPermission 时回退到 SourcePermission
    doc = await db.get(Document, doc_id)
    if doc and doc.source_id:
        sp = await db.execute(
            select(SourcePermission).where(
                SourcePermission.source_id == doc.source_id,
                SourcePermission.user_id == user.id,
            )
        )
        src_perm = sp.scalar_one_or_none()
        if src_perm and _permission_allows(src_perm.permission, required):
            return

    # 最后检查知识库成员权限
    if doc and doc.kb_id:
        from app.models.knowledge_base import KnowledgeBase
        from app.services.kb_permissions import can_edit_kb, can_view_kb

        kb = await db.get(KnowledgeBase, doc.kb_id)
        if kb:
            if required == "write":
                if await can_edit_kb(db, user, kb):
                    return
            else:
                if await can_view_kb(db, user, kb):
                    return

    raise HTTPException(status_code=403, detail=f"无权限{action}该文档")


@router.post("/upload")
async def upload_file(
    request: Request,
    file: UploadFile = File(...),
    kb_id: str | None = None,
    _: None = Depends(upload_rate_limit),
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    suffix = Path(file.filename).suffix.lower()
    if suffix not in ALLOWED_EXT:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {suffix}，允许: {settings.ALLOWED_EXTENSIONS}",
        )

    content = await file.read()
    if len(content) > settings.MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413, detail=f"文件过大，最大 {settings.MAX_FILE_SIZE // 1024 // 1024}MB"
        )

    # 清理文件名防止路径穿越
    upload_dir = Path(settings.UPLOAD_DIR)
    upload_dir.mkdir(parents=True, exist_ok=True)
    safe_filename = Path(file.filename).name
    filepath = upload_dir / safe_filename

    async with aiofiles.open(filepath, "wb") as f:
        await f.write(content)

    if suffix in (".txt", ".md", ".markdown"):
        text_content = content.decode("utf-8", errors="ignore")
    elif suffix == ".pdf":

        def _parse_pdf() -> str:
            import pdfplumber

            try:
                with pdfplumber.open(str(filepath)) as pdf:
                    pages = []
                    for page in pdf.pages:
                        page_text = page.extract_text() or ""
                        for tbl in page.extract_tables() or []:
                            if tbl:
                                rows = ["\t".join(str(c or "") for c in row) for row in tbl]
                                page_text += "\n" + "\n".join(rows)
                        pages.append(page_text)
                    return "\n\n".join(pages)
            except Exception:
                return ""

        text_content = await asyncio.to_thread(_parse_pdf)
    elif suffix == ".docx":

        def _parse_docx() -> str:
            import io

            from docx import Document as DocxDocument

            doc = DocxDocument(io.BytesIO(content))
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

        text_content = await asyncio.to_thread(_parse_docx)
    elif suffix == ".xlsx":

        def _parse_xlsx() -> str:
            import io

            from openpyxl import load_workbook

            wb = load_workbook(io.BytesIO(content), read_only=True)
            sections = []
            for sheet in wb.worksheets:
                sheet_name = sheet.title
                all_rows = list(sheet.iter_rows(values_only=True))
                if not all_rows:
                    continue
                header = [str(c or "") for c in all_rows[0]]
                lines = [f"## 表格：{sheet_name}"]
                for row in all_rows[1:]:
                    pairs = [f"{header[i]}：{str(c)}" for i, c in enumerate(row) if c is not None]
                    if pairs:
                        lines.append("；".join(pairs))
                sections.append("\n".join(lines))
            return "\n\n".join(sections)

        text_content = await asyncio.to_thread(_parse_xlsx)
    else:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {suffix}")

    if not text_content.strip():
        raise HTTPException(status_code=400, detail="文件内容为空")

    result = await db.execute(select(DataSource).where(DataSource.type == "local"))
    source = result.scalars().first()
    if not source:
        source = DataSource(name="本地文件", type="local", created_by=user.id)
        db.add(source)
        await db.flush()

    content_hash = hashlib.md5(text_content.encode()).hexdigest()
    doc = Document(
        source_id=source.id,
        kb_id=kb_id,
        external_id=file.filename,
        title=file.filename,
        stored_filename=safe_filename,
        content=text_content,
        content_hash=content_hash,
    )
    db.add(doc)
    await db.flush()

    # flush 保证 doc.id 可用，上传者默认 write 权限
    db.add(DocumentPermission(document_id=doc.id, user_id=user.id, permission="write"))
    await db.flush()

    from app.tasks.indexing import index_document_task

    index_document_task.delay(str(doc.id))

    from app.services.audit_v2 import record_audit_event

    await record_audit_event(
        db,
        actor_user=user,
        action="document.upload",
        resource_type="document",
        resource_id=doc.id,
        request=request,
        metadata={"title": doc.title, "kb_id": kb_id},
    )

    logger.info(f"文档上传成功: {doc.title} (id={doc.id}, user={user.id})")
    await cache_delete(f"doclist:{user.id}")
    return {"id": str(doc.id), "title": doc.title, "status": "processing"}


@router.get("")
async def list_documents(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    kb_id: str | None = None,
    limit: int = 20,
    offset: int = 0,
):
    if user.role == "admin":
        base = select(Document)
        if kb_id:
            base = base.where(Document.kb_id == kb_id)
    else:
        # JOIN 替代 IN 子查询，性能更优
        base = (
            select(Document)
            .outerjoin(
                DocumentPermission,
                (DocumentPermission.document_id == Document.id)
                & (DocumentPermission.user_id == user.id),
            )
            .outerjoin(
                SourcePermission,
                (SourcePermission.source_id == Document.source_id)
                & (SourcePermission.user_id == user.id),
            )
            .where(
                (DocumentPermission.user_id.isnot(None)) | (SourcePermission.user_id.isnot(None))
            )
            .distinct()
        )
        if kb_id:
            base = base.where(Document.kb_id == kb_id)

    cache_key = f"cache:doclist:{user.id}:{kb_id}:{limit}:{offset}"
    cached = await cache_get(cache_key)
    if cached:
        return cached

    count_query = select(func.count()).select_from(base.subquery())
    total = await db.scalar(count_query)
    result = await db.execute(base.order_by(Document.created_at.desc()).offset(offset).limit(limit))
    docs = result.scalars().all()
    data = {
        "total": total or 0,
        "items": [
            {
                "id": str(d.id),
                "title": d.title,
                "status": d.status or "pending",
                "error_message": d.error_message,
                "retry_count": d.retry_count or 0,
                "indexed_at": str(d.indexed_at) if d.indexed_at else None,
                "kb_id": str(d.kb_id) if d.kb_id else None,
                "created_at": str(d.created_at),
            }
            for d in docs
        ],
    }
    await cache_set(cache_key, data, ttl=30, tags=[f"doclist:{user.id}"])
    return data


@router.get("/stats")
async def document_stats(
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
    kb_id: str | None = None,
):
    # 先获取用户有权限的文档 ID 子查询
    if user.role == "admin":
        perm_q = select(Document.id)
        if kb_id:
            perm_q = perm_q.where(Document.kb_id == kb_id)
    else:
        perm_q = (
            select(Document.id)
            .outerjoin(
                DocumentPermission,
                (DocumentPermission.document_id == Document.id)
                & (DocumentPermission.user_id == user.id),
            )
            .outerjoin(
                SourcePermission,
                (SourcePermission.source_id == Document.source_id)
                & (SourcePermission.user_id == user.id),
            )
            .where(
                (DocumentPermission.user_id.isnot(None)) | (SourcePermission.user_id.isnot(None))
            )
            .distinct()
        )
        if kb_id:
            perm_q = perm_q.where(Document.kb_id == kb_id)

    # 用有权限的文档 ID 子查询做统计
    allowed = perm_q.subquery()
    stats = await db.execute(
        select(
            func.count().label("all"),
            func.sum(case((Document.status == "indexed", 1), else_=0)).label("indexed"),
            func.sum(case((Document.status == "processing", 1), else_=0)).label("processing"),
        )
        .where(Document.id.in_(select(allowed.c.id)))
        .select_from(Document)
    )
    row = stats.one()
    return {
        "all": row.all,
        "indexed": row.indexed,
        "processing": row.processing,
        "others": row.all - row.indexed - row.processing,
    }


@router.get("/{doc_id}")
async def get_document(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    doc = await db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await _require_doc_permission(db, doc_id, user)
    await audit_log(db, str(user.id), "view_doc", "document", doc_id, doc.title)
    return {
        "id": str(doc.id),
        "title": doc.title,
        "content": doc.content[:5000],
        "status": doc.status or "pending",
        "error_message": doc.error_message,
        "retry_count": doc.retry_count or 0,
        "indexed_at": str(doc.indexed_at) if doc.indexed_at else None,
        "created_at": str(doc.created_at),
    }


@router.delete("/{doc_id}")
async def delete_document(
    doc_id: str,
    request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    doc = await db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await _require_doc_permission(db, doc_id, user, "删除", required="write")
    await db.delete(doc)

    from app.services.audit_v2 import record_audit_event

    await record_audit_event(
        db,
        actor_user=user,
        action="document.delete",
        resource_type="document",
        resource_id=doc_id,
        request=request,
        metadata={"title": doc.title},
    )
    await webhook_dispatch(
        db,
        "document.deleted",
        {
            "document_id": doc_id,
            "title": doc.title,
            "deleted_by": str(user.id),
        },
    )
    await cache_delete(f"doclist:{user.id}")
    return {"detail": "已删除"}


@router.get("/{doc_id}/chunks")
async def get_document_chunks(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取文档的全部 chunk（用于前端高亮定位）"""
    doc = await db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await _require_doc_permission(db, doc_id, user)

    result = await db.execute(
        select(DocumentChunk)
        .where(DocumentChunk.document_id == doc_id)
        .order_by(DocumentChunk.chunk_index)
    )
    chunks = result.scalars().all()
    return {
        "document": {"id": str(doc.id), "title": doc.title, "content": doc.content},
        "chunks": [
            {"id": str(c.id), "chunk_index": c.chunk_index, "content": c.content} for c in chunks
        ],
    }


class BatchIds(BaseModel):
    ids: list[str]


async def _filter_authorized_doc_ids(
    data: BatchIds, user: User, db: AsyncSession, required: str = "read"
) -> list[UUID]:
    """验证 UUID 格式并过滤出用户有指定权限的文档 ID

    DocumentPermission 优先；仅当无 DocumentPermission 时才回退到 SourcePermission。
    """
    valid_ids = []
    for doc_id in data.ids:
        try:
            valid_ids.append(UUID(doc_id))
        except ValueError:
            continue

    if not valid_ids or user.role == "admin":
        return valid_ids

    # DocumentPermission
    perm_result = await db.execute(
        select(DocumentPermission.document_id, DocumentPermission.permission).where(
            DocumentPermission.document_id.in_(valid_ids),
            DocumentPermission.user_id == user.id,
        )
    )
    perm_map = {row.document_id: row.permission for row in perm_result}

    # SourcePermission（仅用于无 DocumentPermission 的文档）
    doc_result = await db.execute(
        select(Document.id, Document.source_id).where(Document.id.in_(valid_ids))
    )
    doc_source_map = {row.id: row.source_id for row in doc_result if row.source_id}
    source_ids = set(doc_source_map.values())

    src_perm_map = {}
    if source_ids:
        src_result = await db.execute(
            select(SourcePermission.source_id, SourcePermission.permission).where(
                SourcePermission.source_id.in_(source_ids),
                SourcePermission.user_id == user.id,
            )
        )
        src_perm_map = {row.source_id: row.permission for row in src_result}

    allowed = []
    for uid in valid_ids:
        doc_perm = perm_map.get(uid)
        if doc_perm:
            # DocumentPermission 存在时以其为准
            if _permission_allows(doc_perm, required):
                allowed.append(uid)
        else:
            # 无 DocumentPermission 时回退到 SourcePermission
            src_id = doc_source_map.get(uid)
            if src_id:
                src_perm = src_perm_map.get(src_id)
                if src_perm and _permission_allows(src_perm, required):
                    allowed.append(uid)

    return allowed


@router.post("/batch-delete")
async def batch_delete(
    data: BatchIds,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    valid_ids = await _filter_authorized_doc_ids(data, user, db, required="write")
    if not valid_ids:
        return {"detail": "已删除 0 个文档"}

    await db.execute(delete(Document).where(Document.id.in_(valid_ids)))
    return {"detail": f"已删除 {len(valid_ids)} 个文档"}


@router.post("/batch-reindex")
async def batch_reindex(
    data: BatchIds,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    valid_ids = await _filter_authorized_doc_ids(data, user, db, required="write")
    if not valid_ids:
        return {"detail": "已触发 0 个文档重新索引"}

    from app.tasks.indexing import index_document_task

    for uid in valid_ids:
        # 重置文档状态，与单文档 reindex_kb 行为一致
        doc = await db.get(Document, uid)
        if not doc:
            continue
        doc.status = "pending"
        doc.error_message = None
        await db.flush()
        index_document_task.delay(str(uid))

    return {"detail": f"已触发 {len(valid_ids)} 个文档重新索引"}


@router.post("/{doc_id}/retry-index")
async def retry_index(
    doc_id: str,
    request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """重试文档索引"""
    doc = await db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await _require_doc_permission(db, doc_id, user, "重试索引", required="write")

    if doc.status not in ("failed", "pending", "processing"):
        raise HTTPException(status_code=400, detail=f"当前状态 '{doc.status}' 不允许重试")

    doc.retry_count = (doc.retry_count or 0) + 1
    doc.error_message = None
    doc.status = "pending"
    await db.commit()

    # 清除可能残留的 Redis 锁，确保重试任务不被阻塞
    from app.tasks.indexing import clear_index_lock

    lock_cleared = await clear_index_lock(doc_id)
    if not lock_cleared:
        logger.warning(f"文档 {doc_id} Redis 锁清除失败或锁不存在，worker 将跳过索引")

    from app.tasks.indexing import index_document_task

    index_document_task.delay(str(doc.id))

    from app.services.audit_v2 import record_audit_event

    await record_audit_event(
        db,
        actor_user=user,
        action="document.retry_index",
        resource_type="document",
        resource_id=doc_id,
        request=request,
        metadata={"title": doc.title, "retry_count": doc.retry_count},
    )

    await cache_delete(f"doclist:{user.id}")
    return {
        "id": str(doc.id),
        "title": doc.title,
        "status": doc.status,
        "error_message": doc.error_message,
        "retry_count": doc.retry_count,
        "indexed_at": str(doc.indexed_at) if doc.indexed_at else None,
    }


@router.get("/{doc_id}/preview")
async def preview_document(
    doc_id: str,
    request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """文档预览：txt/md 返回文本内容，pdf/docx/xlsx 返回 download_only"""
    doc = await db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await _require_doc_permission(db, doc_id, user, "预览")

    # 推断文件类型
    title_lower = (doc.title or "").lower()
    if title_lower.endswith((".txt", ".md", ".markdown")):
        file_type = "txt" if title_lower.endswith(".txt") else "md"
        preview_mode = "text"
    elif title_lower.endswith(".pdf"):
        file_type = "pdf"
        preview_mode = "download_only"
    elif title_lower.endswith(".docx"):
        file_type = "docx"
        preview_mode = "download_only"
    elif title_lower.endswith(".xlsx"):
        file_type = "xlsx"
        preview_mode = "download_only"
    else:
        file_type = "unknown"
        preview_mode = "download_only"

    result: dict = {
        "document_id": str(doc.id),
        "title": doc.title,
        "file_type": file_type,
        "status": doc.status or "pending",
        "preview_mode": preview_mode,
        "download_url": f"/api/v1/documents/{doc.id}/file",
    }

    if preview_mode == "text" and doc.content:
        result["content"] = doc.content[:20000]

    from app.services.audit_v2 import record_audit_event

    await record_audit_event(
        db,
        actor_user=user,
        action="document.preview",
        resource_type="document",
        resource_id=doc_id,
        request=request,
        metadata={"document_id": doc_id, "file_type": file_type},
    )

    return result


@router.get("/{doc_id}/chunks/{chunk_id}/locator")
async def get_chunk_locator(
    doc_id: str,
    chunk_id: str,
    request: Request,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """获取 chunk 定位信息（page/section/locator）"""
    doc = await db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await _require_doc_permission(db, doc_id, user, "查看")

    chunk = await db.get(DocumentChunk, chunk_id)
    if not chunk or str(chunk.document_id) != doc_id:
        raise HTTPException(status_code=404, detail="文档块不存在或不属于该文档")

    meta = chunk.metadata_ or {}
    page = meta.get("page") or meta.get("page_number")
    section = meta.get("section") or meta.get("heading")

    locator: dict
    if page is not None:
        locator = {"type": "page", "value": str(page)}
    elif section:
        locator = {"type": "text", "value": str(section)}
    else:
        locator = {"type": "chunk", "value": str(chunk.id)}

    from app.services.audit_v2 import record_audit_event

    await record_audit_event(
        db,
        actor_user=user,
        action="document.locator_view",
        resource_type="document",
        resource_id=doc_id,
        request=request,
        metadata={"document_id": doc_id, "chunk_id": chunk_id, "locator_type": locator["type"]},
    )

    result: dict = {
        "document_id": str(doc.id),
        "chunk_id": str(chunk.id),
        "snippet": chunk.content[:300],
        "locator": locator,
    }
    if page is not None:
        result["page"] = int(page) if isinstance(page, (int, float, str)) and str(page).isdigit() else page
    if section:
        result["section"] = str(section)
    return result


@router.get("/{doc_id}/file")
async def download_file(
    doc_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """返回文档原始文件"""
    doc = await db.get(Document, doc_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await _require_doc_permission(db, doc_id, user)

    filepath = Path(settings.UPLOAD_DIR) / (doc.stored_filename or Path(doc.title).name)
    if not filepath.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    await audit_log(db, str(user.id), "download_file", "document", doc_id, doc.title)
    return FileResponse(str(filepath), filename=doc.title)


@router.get("/chunks/{chunk_id}")
async def get_chunk_detail(
    chunk_id: str,
    user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """通过 chunk_id 获取 chunk 所属文档及全部 chunk"""
    chunk = await db.get(DocumentChunk, chunk_id)
    if not chunk:
        raise HTTPException(status_code=404, detail="文档块不存在")

    doc = await db.get(Document, chunk.document_id)
    if not doc:
        raise HTTPException(status_code=404, detail="文档不存在")
    await _require_doc_permission(db, str(doc.id), user)

    result = await db.execute(
        select(DocumentChunk)
        .where(DocumentChunk.document_id == doc.id)
        .order_by(DocumentChunk.chunk_index)
    )
    all_chunks = result.scalars().all()
    return {
        "document": {"id": str(doc.id), "title": doc.title, "content": doc.content},
        "highlight_chunk_id": str(chunk_id),
        "chunks": [
            {"id": str(c.id), "chunk_index": c.chunk_index, "content": c.content}
            for c in all_chunks
        ],
    }
